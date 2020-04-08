import os
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import math
import seaborn as sn
from numpy import log as ln
from docx import Document
from docx.shared import Pt

class Con:
    def __init__(self, dataframe):
        self.__dataframe = dataframe
        self.__datalist = dataframe.values.tolist()

    @property
    def dataframe(self):
        return self.__dataframe

    @property
    def datalist(self):
        return self.__datalist

    def clearlist(self):
        return sorted([val for val in self.dataframe if not np.isnan(val)])

    def length(self):
        return len(self.dataframe)

    def missingpercent(self):
        return sum(np.isnan(val) for val in self.dataframe) / self.length() * 100
    
    def cardinality(self):     
        return len(set(self.clearlist()))

    def max(self):
        return max(self.clearlist())

    def min(self):
        return min(self.clearlist())

    def firstquarter(self):
        index = len(self.clearlist()) % 4
        quarter = int(len(self.clearlist()) / 4)  
        if index == 0 or index == 1:
            return (self.clearlist()[quarter] + self.clearlist()[quarter - 1]) / 2 
        elif index == 2 or index == 3:
            return self.clearlist()[quarter]

    def thirdquarter(self):
        index = len(self.clearlist()) * 3 % 4
        quarter = int(len(self.clearlist()) * 3 / 4)
        if index == 0:
            return (self.clearlist()[quarter] + self.clearlist()[quarter -1]) / 2
        elif index == 3:
            return (self.clearlist()[quarter] + self.clearlist()[quarter + 1]) / 2
        elif index == 2 or index == 1:
            return self.clearlist()[quarter]

    def avg(self):
        return sum(self.clearlist()) / len(self.clearlist())

    def median(self):
        index = len(self.clearlist()) % 2
        middle = int(len(self.clearlist()) / 2)
        return self.clearlist()[middle] if index == 0 else (self.clearlist()[middle] + self.clearlist()[middle + 1]) / 2

    def variance(self):
        avg = int(self.avg()) 
        return sum([(val - avg) ** 2 for val in self.clearlist()]) / len(self.clearlist())

    def standartdev(self):
        return math.sqrt(self.variance())

class Categoric:
    def __init__(self, dataframe):
        self.__dataframe = dataframe
        self.__datalist = dataframe.values.tolist()

    @property
    def dataframe(self):
        return self.__dataframe

    @property
    def datalist(self):
        return self.__datalist

    def clearlist(self):
        return [val for val in self.datalist if val is not np.nan]

    def length(self):
        return len(self.dataframe)

    def missingpercent(self):
        return sum( 1 for val in self.dataframe if val is np.nan) / self.length() * 100
    
    def cardinality(self):     
        return len(set(self.clearlist()))

    def valueoccurences(self):
        alldata = self.clearlist()
        dictionary = {val:alldata.count(val) for val in set(self.clearlist())}
        return {k: v for k, v in sorted(dictionary.items(), key=lambda item: item[1])}
    
    def mode(self):
        return list(self.valueoccurences().keys())[-1]
    
    def modecount(self):
        return list(self.valueoccurences().values())[-1]

    def modepercentage(self):
        return self.modecount() / self.length() * 100

    def mode2(self):
        return list(self.valueoccurences().keys())[-2]
    
    def mode2count(self):
        return list(self.valueoccurences().values())[-2]

    def mode2percentage(self):
        return self.mode2count() / self.length() * 100

def read(file):
    data = pd.read_csv(file)
    return data

def sepheaders(data):
    headers = data.columns   
    cont = []
    categ = []
    for header in headers:
        if data[header].dtype == np.float64:
            cont.append(header)
        else:
            categ.append(header)
    return cont, categ

def drawhistogram(data):
    cont, categ = sepheaders(data)
    for header in cont:
        plt.figure(num=header)
        binvalue =int(1 + 3.22 * ln(len(data[header])))
        data[header].hist(bins=binvalue)
        plt.xlabel('Values')
        plt.ylabel('Frequency')
        plt.title(header+" histogram")
    plt.show()

def drawbarplot(data, header):
    data[header].plot.bar(rot=0)

def drawbarplotcategoric(data):
    cont, categ = sepheaders(data)
    for header in categ:
        plt.figure(num=header)
        data[header].value_counts().plot(kind='bar')
        plt.xlabel('Values')
        plt.ylabel('Frequency')
        plt.title(header+" histogram")
    plt.show()

def drawscatter(data, header, header1):
    plt.figure(num=header+header1)
    plt.scatter(data[header], data[header1])
    plt.xlabel(header)
    plt.ylabel(header1)
    plt.show()

def splom(data):
    cont, categ = sepheaders(data)
    pd.plotting.scatter_matrix(data, alpha=0.2)
    plt.show()

def drawboxplot(data, header):
    df = pd.DataFrame()

    # data[header].plot.box()
    # data[header].plot.box(grid='True')
    # plt.show()


def drawcorrelation(data):
    corrMatrix = data.corr()
    sn.heatmap(corrMatrix, annot=True)
    plt.show()

def todoc(dat, continueslist, categorylist):
    document = Document()
    
    tstyle = document.styles['Normal']
    font = tstyle.font
    font.name = "Times New Roman"
    font.size = Pt(9)
    
    table = document.add_table(rows=1,cols=11)
    table.style = document.styles['Table Grid']
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pavadinimas'
    hdr_cells[1].text = 'Bendras reikšmių skaičius'
    hdr_cells[2].text = 'Trūkstamų skaičius'
    hdr_cells[3].text = 'Kardinalumas'
    hdr_cells[4].text = 'Minimali reikšmė'
    hdr_cells[5].text = 'Maksimali reikšmė'
    hdr_cells[6].text = '1-asis kvartilis'
    hdr_cells[7].text = '2-asis kvartilis'
    hdr_cells[8].text = 'Vidurkis'
    hdr_cells[9].text = 'Mediana'
    hdr_cells[10].text = 'Standartinis nuokrypis'
    
    for header in continueslist:
        contobject = Con(dat[header])
        row_cells = table.add_row().cells
        row_cells[0].text = header
        row_cells[1].text = str(contobject.length())
        row_cells[2].text = str(round(contobject.missingpercent(),3)) + "%"
        row_cells[3].text = str(contobject.cardinality())
        row_cells[4].text = str(contobject.min())
        row_cells[5].text = str(contobject.max())
        row_cells[6].text = str(contobject.firstquarter())
        row_cells[7].text = str(contobject.thirdquarter())
        row_cells[8].text = str(round(contobject.avg(),3))
        row_cells[9].text = str(contobject.median())
        row_cells[10].text = str(round(contobject.standartdev(),3))
    
    document.add_page_break()

    table1 = document.add_table(rows=1, cols=10)
    table1.style = document.styles['Table Grid']

    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = 'Pavadinimas'
    hdr1_cells[1].text = 'Bendras reikšmių skaičius'
    hdr1_cells[2].text = 'Trūkstamų skaičius'
    hdr1_cells[3].text = 'Kardinalumas'
    hdr1_cells[4].text = 'Moda'
    hdr1_cells[5].text = 'Modos dažnumas'
    hdr1_cells[6].text = 'Modos dažnumas procentais'
    hdr1_cells[7].text = '2-oji moda'
    hdr1_cells[8].text = '2-osios modos dažnumas'
    hdr1_cells[9].text = '2-osios modos dažnumas procentais'

    for header in categorylist:
        catobject = Categoric(dat[header])
        row_cells = table1.add_row().cells
        row_cells[0].text = header
        row_cells[1].text = str(catobject.length())
        row_cells[2].text = str(round(catobject.missingpercent(), 3)) + "%"
        row_cells[3].text = str(catobject.cardinality())
        row_cells[4].text = str(catobject.mode())
        row_cells[5].text = str(catobject.modecount())
        row_cells[6].text = str(round(catobject.modepercentage(), 3)) + "%"
        row_cells[7].text = str(catobject.mode2())
        row_cells[8].text = str(catobject.mode2count())
        row_cells[9].text = str(round(catobject.mode2percentage(), 3)) + "%"
    
    document.save('./App/data/data.docx')

def results(data):
    cont, categ = sepheaders(data)
    # drawhistogram(data)
    todoc(data, cont, categ)

 # Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope Process
def main():
    dat = read("./App/Data/Game_Data.csv")
    cont = Con(dat["Year_of_Release"])
    # print(dat["Year_of_Release"].describe())
    # print(cont.thirdquarter())

    # results(dat)
    # drawhistogram(dat)
    # drawbarplotcategoric(dat)
    drawscatter(dat, "Other_Sales", "NA_Sales")
    # splom(dat)
    # drawboxplot(dat,"Platform")
    # drawcorrelation(dat)

if __name__ == "__main__":
    main()
