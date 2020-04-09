import os
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import math
import seaborn as sn
from docx import Document
from docx.shared import Pt
from Code.continues import Con
from Code.categoric import Categoric

def read(file):
    data = pd.read_csv(file)
    return data

def todoc(dat, continueslist, categorylist):
    document = Document()
    
    tstyle = document.styles['Normal']
    font = tstyle.font
    font.name = "Times New Roman"
    font.size = Pt(9)
    
    table = document.add_table(rows=1,cols=11)
    table.style = document.styles['Table Grid']
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pavadinimas'
    hdr_cells[1].text = 'Bendras reikšmių skaičius'
    hdr_cells[2].text = 'Trūkstamų skaičius'
    hdr_cells[3].text = 'Kardinalumas'
    hdr_cells[4].text = 'Minimali reikšmė'
    hdr_cells[5].text = 'Maksimali reikšmė'
    hdr_cells[6].text = '1-asis kvartilis'
    hdr_cells[7].text = '2-asis kvartilis'
    hdr_cells[8].text = 'Vidurkis'
    hdr_cells[9].text = 'Mediana'
    hdr_cells[10].text = 'Standartinis nuokrypis'
    
    for header in continueslist:
        contobject = Con(dat[header])
        row_cells = table.add_row().cells
        row_cells[0].text = header
        row_cells[1].text = str(contobject.length())
        row_cells[2].text = str(round(contobject.missingpercent(),3)) + "%"
        row_cells[3].text = str(contobject.cardinality())
        row_cells[4].text = str(contobject.min())
        row_cells[5].text = str(contobject.max())
        row_cells[6].text = str(contobject.firstquarter())
        row_cells[7].text = str(contobject.thirdquarter())
        row_cells[8].text = str(round(contobject.avg(),3))
        row_cells[9].text = str(contobject.median())
        row_cells[10].text = str(round(contobject.standartdev(),3))
    
    document.add_page_break()

    table1 = document.add_table(rows=1, cols=10)
    table1.style = document.styles['Table Grid']

    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = 'Pavadinimas'
    hdr1_cells[1].text = 'Bendras reikšmių skaičius'
    hdr1_cells[2].text = 'Trūkstamų skaičius'
    hdr1_cells[3].text = 'Kardinalumas'
    hdr1_cells[4].text = 'Moda'
    hdr1_cells[5].text = 'Modos dažnumas'
    hdr1_cells[6].text = 'Modos dažnumas procentais'
    hdr1_cells[7].text = '2-oji moda'
    hdr1_cells[8].text = '2-osios modos dažnumas'
    hdr1_cells[9].text = '2-osios modos dažnumas procentais'

    for header in categorylist:
        catobject = Categoric(dat[header])
        row_cells = table1.add_row().cells
        row_cells[0].text = header
        row_cells[1].text = str(catobject.length())
        row_cells[2].text = str(round(catobject.missingpercent(), 3)) + "%"
        row_cells[3].text = str(catobject.cardinality())
        row_cells[4].text = str(catobject.mode())
        row_cells[5].text = str(catobject.modecount())
        row_cells[6].text = str(round(catobject.modepercentage(), 3)) + "%"
        row_cells[7].text = str(catobject.mode2())
        row_cells[8].text = str(catobject.mode2count())
        row_cells[9].text = str(round(catobject.mode2percentage(), 3)) + "%"
    
    document.save('./Lab1/data/data.docx')